from  flask import Flask, request, render_template
import math
import io
import os

try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    from sentence_transformers import SentenceTransformer, util
    _HAVE_SENT_TRANS = True
except Exception:
    SentenceTransformer = None
    util = None
    _HAVE_SENT_TRANS = False

app = Flask(__name__)

app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Model wrapper: use sentence-transformers when available, otherwise a lightweight dummy
class ModelWrapper:
    def __init__(self):
        if _HAVE_SENT_TRANS and SentenceTransformer is not None:
            try:
                self.model = SentenceTransformer('all-MiniLM-L6-v2')
                self.use_real = True
            except Exception:
                self.model = None
                self.use_real = False
        else:
            self.model = None
            self.use_real = False

    def encode(self, inputs):
        # Accept both single string and list/iterable
        if isinstance(inputs, str):
            inputs = [inputs]

        if self.use_real and self.model is not None:
            return self.model.encode(list(inputs))

        # Simple fallback: character-frequency vector (a-z)
        def _char_vec(s):
            s = (s or '').lower()
            vec = [0.0] * 26
            for ch in s:
                idx = ord(ch) - ord('a')
                if 0 <= idx < 26:
                    vec[idx] += 1.0
            # normalize
            norm = math.sqrt(sum(v * v for v in vec))
            if norm > 0:
                vec = [v / norm for v in vec]
            return vec

        return [_char_vec(s) for s in inputs]

    @staticmethod
    def cos_sim(a, b):
        # if using real model, util.cos_sim will be called by caller; this is for fallback
        # a and b are sequences of numbers
        if hasattr(a, 'tolist'):
            try:
                a = a.tolist()
            except Exception:
                pass
        if hasattr(b, 'tolist'):
            try:
                b = b.tolist()
            except Exception:
                pass
        dot = sum(x * y for x, y in zip(a, b))
        na = math.sqrt(sum(x * x for x in a))
        nb = math.sqrt(sum(y * y for y in b))
        if na == 0 or nb == 0:
            return 0.0
        return dot / (na * nb)

# instantiate model wrapper
model_wrapper = ModelWrapper()

#Allowed file extensions
ALLOWED_EXTENSIONS = {'pdf', 'docx'}

def allowed_file(filename):
    return '.' in filename and \
        filename.rsplit('.',1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(pdf_file):

    """
        Extract text from a PDF file
    """

    text = ""

    try:
        pdf_bytes = pdf_file.read()
        if fitz is not None:
            # PyMuPDF path
            with fitz.open(stream=pdf_bytes, filetype='pdf') as doc:
                for page in doc:
                    text += page.get_text()
        elif PyPDF2 is not None:
            # Fallback to PyPDF2
            pdf_stream = io.BytesIO(pdf_bytes)
            reader = PyPDF2.PdfReader(pdf_stream)
            for page in reader.pages:
                try:
                    text += page.extract_text() or ''
                except Exception:
                    continue
        else:
            raise RuntimeError('No PDF parsing library available (fitz or PyPDF2)')
    except Exception as e:
        print(f"Error reading PDF: {e}")
        text = f"Error extracting text from PDF: {e}"
    
    return text

def extract_text_from_docx(docx_file):

    """
        Extract text from a DOCX file
    """

    text = ""

    try:
        if Document is None:
            raise RuntimeError('python-docx is not available')
        # python-docx accepts a file-like object
        document = Document(io.BytesIO(docx_file.read())) if not isinstance(docx_file, str) else Document(docx_file)
        for paragraph in document.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        text = f"Error extracting text from DOCX: {e}"
    
    return text

def analysis_resume(resume_text, job_description):
    "performs simple semantic matching between resume and comma-separated skills"

    required_skills = [skill.strip().lower() for skill in job_description.split(',') if skill.strip()]

    if not required_skills:
        return {'matched_skills': [], 'missing_skills': [], 'match_score': 0.0}

    # get embeddings
    req_emb = model_wrapper.encode(required_skills)
    resume_emb = model_wrapper.encode(resume_text)[0]

    matched_skills = []
    missing_skills = []

    for i, skill in enumerate(required_skills):
        skill_embedding = req_emb[i]
        if _HAVE_SENT_TRANS and util is not None and model_wrapper.use_real:
            # when using real model, util.cos_sim returns a tensor
            try:
                similarity_score = util.cos_sim(resume_emb, skill_embedding).item()
            except Exception:
                # fallback to wrapper's cos_sim
                similarity_score = model_wrapper.cos_sim(resume_emb, skill_embedding)
        else:
            similarity_score = model_wrapper.cos_sim(resume_emb, skill_embedding)

        entry = {'skill': skill, 'score': float(similarity_score)}
        if similarity_score > 0.35:
            matched_skills.append(entry)
        else:
            missing_skills.append(entry)

    match_score = len(matched_skills) / len(required_skills) if required_skills else 0.0

    return {
        'matched_skills': matched_skills,
        'missing_skills': missing_skills,
        'match_score': match_score,
    }

@app.route('/', methods = ['GET', 'POST'])
def upload_and_process():
    extracted_text = None

    analysis_results = None

    if request.method == 'POST':

        #check if post request has file part
        if 'resume' not in request.files:
            return render_template('index.html', error = 'No file part')
        
        resume_file = request.files.get('resume')

        job_description = request.form.get('job_description', '')

        #if jd is empty
        if job_description == '':
            return render_template('index.html', error = 'No job description')
        
        #if user does not select file, browser also submits an empty part without filename
        if resume_file.filename == '':
            return render_template('index.html', error = 'No selected file')

        # process the uploaded file
        if resume_file and allowed_file(resume_file.filename):
            file_extension = resume_file.filename.rsplit('.', 1)[1].lower()
            file_stream = io.BytesIO(resume_file.read())

            if file_extension == 'pdf':
                extracted_text = extract_text_from_pdf(file_stream)
            elif file_extension == 'docx':
                extracted_text = extract_text_from_docx(file_stream)
            else:
                extracted_text = "Unsupported file type"

            if extracted_text and job_description:
                analysis_results = analysis_resume(extracted_text, job_description)

    return render_template('index.html', extracted_text=extracted_text, analysis_results=analysis_results)

    

if __name__ == '__main__':
    # If run with argument 'test', run a quick CLI test and exit. Otherwise start the server.
    import sys
    if len(sys.argv) > 1 and sys.argv[1] == 'test':
        sample_resume = 'Experienced Python developer. Worked with Flask, Django and REST APIs.'
        sample_jd = 'python, flask, docker'
        print('Resume:', sample_resume)
        print('Job description:', sample_jd)
        result = analysis_resume(sample_resume, sample_jd)
        print('\nAnalysis result:')
        print(result)
    else:
        app.run(debug=True)